"""One-shot upload readiness check for FINAL_SUBMISSION_PACKAGE."""
from pathlib import Path
from docx import Document

root = Path(__file__).resolve().parents[1] / "FINAL_SUBMISSION_PACKAGE"
checks = []

t = Document(root / "01_Title_Page.docx")
tb = "\n".join(p.text for p in t.paragraphs)
checks.append(("Title matches locked title", "Do Subsidised Solar Pumps Really Cut Carbon" in tb))
checks.append(("Title has Harsh ORCID", "0009-0008-7394-130X" in tb))
checks.append(("Title has Gunjan ORCID", "0000-0001-6004-7642" in tb))
checks.append(("Title has CRediT", "Conceptualisation" in tb or "Conceptualization" in tb))
checks.append(("Title has funding + competing", "competing interests" in tb.lower() and "Funding" in tb))

h = Document(root / "02_Highlights.docx")
hs = [
    p.text for p in h.paragraphs
    if p.text and not p.text.startswith("Highlights") and not p.text.startswith("Energy Policy")
]
checks.append(("5 highlights", len(hs) == 5))
checks.append(("all highlights max 85 chars", all(len(x) <= 85 for x in hs)))
checks.append(("highlight uses ~31%", any("31%" in x for x in hs)))

m = Document(root / "03_Manuscript.docx")
mb = "\n".join(p.text for p in m.paragraphs)
nimg = sum(1 for r in m.part.rels.values() if "image" in r.reltype)
checks.append(("MS anonymised (no Harsh Dagar)", "Harsh Dagar" not in mb))
checks.append(("MS anonymised (no ORCID)", "orcid" not in mb.lower()))
checks.append(("MS has 4 embedded figures", nimg == 4))
checks.append(("MS DOI filled (not placeholder)", "DOI to be added" not in mb))
checks.append(("MS AI declaration filled (no stub)", "NAME OF TOOL" not in mb))
checks.append(("MS section 6 rewritten", "What the two results mean together" in mb))

c = Document(root / "04_Cover_Letter.docx")
cb = "\n".join(p.text for p in c.paragraphs)
checks.append(("Cover addressed to Energy Policy", "Energy Policy" in cb))
checks.append(("Cover names both authors", "Gunjan Bhandari" in cb and "Harsh Dagar" in cb))
checks.append(("Cover has 1.93 / 1.65 / 31%", "1.93" in cb and "1.65" in cb and "31%" in cb))

s = Document(root / "05_Supplementary_Information.docx")
sb = "\n".join(p.text for p in s.paragraphs)
simg = sum(1 for r in s.part.rels.values() if "image" in r.reltype)
checks.append(("SI has 10 figures", simg == 10))
checks.append(("SI labels Tables S1-S3", "Table S1" in sb and "Table S2" in sb and "Table S3" in sb))
checks.append(("SI has authors", "Harsh Dagar" in sb and "Gunjan Bhandari" in sb))
checks.append(("SI intro has no raw paths", "outputs/tables" not in sb))

figs = list((root / "figures").glob("V*.png"))
tabs = list((root / "tables").glob("SI_S*.csv"))
checks.append(("figures/ has 10 PNGs", len(figs) == 10))
checks.append(("tables/ has 3 SI CSVs", len(tabs) == 3))

print("UPLOAD READINESS CHECK")
ok = 0
blockers = []
for name, val in checks:
    mark = "PASS" if val else "NOTE"
    if val:
        ok += 1
    else:
        blockers.append(name)
    print(f"  [{mark}] {name}")
print(f"\n{ok}/{len(checks)} checks true as written")
print("\nHighlights:")
for x in hs:
    print(f"  ({len(x)}) {x}")
if blockers:
    print("\nItems flagged (may be blockers or expected notes):")
    for b in blockers:
        print(f"  - {b}")
