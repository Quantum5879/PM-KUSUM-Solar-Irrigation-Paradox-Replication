"""Test: extract OMML from a pandoc-generated docx and re-inject it into a
python-docx-authored document, to confirm the round-trip produces a valid,
openable docx with native math objects."""
import subprocess
import tempfile
import zipfile
from pathlib import Path

from lxml import etree
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NSMAP = {"m": M_NS}

MATHML_NSDECLS = (
    'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
)


def latex_to_omml_xml(latex: str, display: bool) -> str:
    delim = "$$" if display else "$"
    md_text = f"{delim}{latex}{delim}"
    with tempfile.TemporaryDirectory() as td:
        md_path = Path(td) / "eq.md"
        docx_path = Path(td) / "eq.docx"
        md_path.write_text(md_text, encoding="utf-8")
        r = subprocess.run(
            ["pandoc", str(md_path), "-o", str(docx_path)],
            capture_output=True, text=True, timeout=15,
        )
        if r.returncode != 0:
            raise RuntimeError(f"pandoc failed: {r.stderr}")
        with zipfile.ZipFile(docx_path) as z:
            xml_bytes = z.read("word/document.xml")
    root = etree.fromstring(xml_bytes)
    if display:
        node = root.find(".//m:oMathPara", NSMAP)
    else:
        node = root.find(".//m:oMath", NSMAP)
    if node is None:
        raise ValueError(f"No math node found for: {latex}")
    return etree.tostring(node, encoding="unicode")


def add_block_equation(doc: Document, latex: str):
    xml_str = latex_to_omml_xml(latex, display=True)
    p = doc.add_paragraph()
    element = parse_xml(f"<m:tmp {MATHML_NSDECLS}>{xml_str}</m:tmp>")
    # element[0] is the actual oMathPara node; move its children into the paragraph
    omath_para = element[0]
    p._p.append(omath_para)
    return p


def add_inline_equation(paragraph, latex: str):
    xml_str = latex_to_omml_xml(latex, display=False)
    element = parse_xml(f"<m:tmp {MATHML_NSDECLS}>{xml_str}</m:tmp>")
    omath = element[0]
    paragraph._p.append(omath)


def main():
    doc = Document()
    doc.add_heading("Equation OMML round-trip test", level=1)

    add_block_equation(doc, r"g(Y_{st}) = \beta \cdot \text{KUSUM}_{st} + \gamma \log(\text{GDPpc})_{st} + \alpha_s + \lambda_t + \varepsilon_{st}")

    p = doc.add_paragraph()
    run = p.add_run("State fixed effects (")
    add_inline_equation(p, r"\alpha_s")
    run2 = p.add_run(") absorb time-invariant characteristics, and year fixed effects (")
    add_inline_equation(p, r"\lambda_t")
    run3 = p.add_run(") absorb national shocks.")

    out = Path(__file__).resolve().parent / "_eq_test3_output.docx"
    doc.save(out)
    print(f"Saved {out}")

    # Validate: reopen with python-docx and check it doesn't error, and confirm oMath present in XML
    doc2 = Document(str(out))
    print("Reopened OK. Paragraph count:", len(doc2.paragraphs))
    with zipfile.ZipFile(out) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    print("Contains oMath:", "<m:oMath" in xml)
    print("Contains oMathPara:", "<m:oMathPara" in xml)


if __name__ == "__main__":
    main()
