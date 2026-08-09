"""
Energy Policy submission pack (separate files, as Elsevier Editorial Manager expects):

  FINAL_SUBMISSION_PACKAGE/
    00_UPLOAD_THESE.txt
    01_Title_Page.docx          - authors, affiliations, ORCID, CRediT, declarations
    02_Highlights.docx          - 3-5 bullets, each <=85 characters
    03_Manuscript.docx          - NO author details; double-spaced; continuous line numbers
    04_Cover_Letter.docx        - short submission letter

Source manuscript markdown remains the master text; identity is stripped from 03 only.
"""
from __future__ import annotations

import re
import subprocess
import tempfile
import zipfile
from functools import lru_cache
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "outputs" / "manuscript_PAPER3_NatureSustainability.md"
OUTDIR = ROOT / "FINAL_SUBMISSION_PACKAGE"
OUTDIR.mkdir(parents=True, exist_ok=True)

BODY_SIZE = 12
TABLE_SIZE = 10

TITLE = (
    "Do Subsidised Solar Pumps Really Cut Carbon or Shift the Cost to Groundwater? "
    "Evidence from India's PM-KUSUM"
)

HIGHLIGHTS = [
    "KUSUM intensity raises GW stage by +0.10 pp (IV; F = 20.6).",
    "Component B drives the effect (F = 54.3); Component C is weak.",
    "Carbon rebound offsets ~31% of gains in over-exploited zones.",
    "IV sign survives dropping Punjab, Haryana and Rajasthan.",
    "Condition Component B and haircut credits; do not halt pumps.",
]

# Sections that reveal authorship, kept on the title page only
STRIP_SECTIONS = {
    "ACKNOWLEDGEMENTS",
    "AUTHOR CONTRIBUTIONS",
    "COMPETING INTERESTS",
}

UNNUMBERED = {
    "ABSTRACT", "METHODS SUMMARY", "DATA AVAILABILITY", "CODE AVAILABILITY",
    "SUPPLEMENTARY INFORMATION", "REFERENCES", "HIGHLIGHTS",
    "DECLARATION OF GENERATIVE AI AND AI-ASSISTED TECHNOLOGIES IN THE MANUSCRIPT PREPARATION PROCESS",
}


def set_run_font(run, name="Times New Roman", size=BODY_SIZE, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_double(paragraph):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)


def enable_line_numbers(section) -> None:
    sect_pr = section._sectPr
    for child in list(sect_pr):
        if child.tag == qn("w:lnNumType"):
            sect_pr.remove(child)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "continuous")
    sect_pr.append(ln)


def add_page_number(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run()
    set_run_font(run, size=10)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def force_heading_font(paragraph, size: int):
    """python-docx Heading styles often ignore theme fonts; stamp TNR on every run."""
    set_double(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=True)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def configure_styles(doc: Document, *, line_numbers: bool = False):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(BODY_SIZE)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for i, size in [(1, 14), (2, 12), (3, 12)]:
        style = doc.styles[f"Heading {i}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        add_page_number(section)
        if line_numbers:
            enable_line_numbers(section)


def add_para(doc, text="", *, size=BODY_SIZE, bold=False, italic=False, center=False, indent=False, justify=True):
    p = doc.add_paragraph()
    set_double(p)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent and text and not text.startswith("**"):
        p.paragraph_format.first_line_indent = Cm(1.27)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_figure(doc: Document, stem: str, caption: str, width_in: float = 6.3):
    fig_dir = ROOT / "outputs" / "figures" / "executive"
    png = fig_dir / f"{stem}.png"
    p = doc.add_paragraph()
    set_double(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    if png.exists():
        run = p.add_run()
        run.add_picture(str(png), width=Inches(width_in))
    else:
        run = p.add_run(f"[Missing figure: {stem}.png]")
        set_run_font(run, size=10, italic=True)

    cap = doc.add_paragraph()
    set_double(cap)
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(6)
    run = cap.add_run(caption.strip())
    set_run_font(run, size=10, italic=True)
    return cap


def add_formatted_runs(paragraph, text: str, base_size=BODY_SIZE):
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Courier New", size=base_size - 1)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=base_size, italic=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "666666")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.autofit = True
    for i, row_data in enumerate(rows):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            val = row_data[j] if j < len(row_data) else ""
            val = re.sub(r"\*\*([^*]+)\*\*", r"\1", val)
            run = p.add_run(val.strip())
            set_run_font(run, size=TABLE_SIZE, bold=(i == 0))
            set_cell_border(cell)
        if i == 0:
            for cell in table.rows[0].cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "E7E6E6")
                shd.set(qn("w:val"), "clear")
                tcPr.append(shd)
    add_para(doc, "")


def parse_md_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.match(r"^\|[\s:\-|]+\|$", line):
            i += 1
            continue
        rows.append([c.strip() for c in line.strip("|").split("|")])
        i += 1
    return rows, i


def is_equation(line: str) -> bool:
    s = line.strip()
    return s.startswith("$$") or (s.startswith("$") and s.endswith("$") and len(s) > 2)


# --------------------------------------------------------------------------- LaTeX -> native Word math (OMML)
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_MATH_NSDECLS = (
    'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
)


@lru_cache(maxsize=None)
def _latex_to_omml_xml(latex: str, display: bool) -> str:
    """Convert one LaTeX snippet to an OMML XML string via pandoc, so the equation
    becomes a native, in-editor Word math object instead of Unicode-substituted text.
    Cached because the same equation string can recur across the manuscript."""
    delim = "$$" if display else "$"
    md_text = f"{delim}{latex}{delim}"
    with tempfile.TemporaryDirectory() as td:
        md_path = Path(td) / "eq.md"
        docx_path = Path(td) / "eq.docx"
        md_path.write_text(md_text, encoding="utf-8")
        result = subprocess.run(
            ["pandoc", str(md_path), "-o", str(docx_path)],
            capture_output=True, text=True, timeout=20,
        )
        if result.returncode != 0:
            raise RuntimeError(f"pandoc failed on {latex!r}: {result.stderr}")
        with zipfile.ZipFile(docx_path) as z:
            xml_bytes = z.read("word/document.xml")
    root = etree.fromstring(xml_bytes)
    tag = "oMathPara" if display else "oMath"
    node = root.find(f".//{{{M_NS}}}{tag}")
    if node is None:
        raise ValueError(f"pandoc produced no {tag} for: {latex!r}")
    return etree.tostring(node, encoding="unicode")


def add_block_equation_native(doc: Document, latex: str):
    """Append a centred, native OMML display equation as its own paragraph."""
    xml_str = _latex_to_omml_xml(latex, display=True)
    p = doc.add_paragraph()
    set_double(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wrapper = parse_xml(f"<m:tmp {_MATH_NSDECLS}>{xml_str}</m:tmp>")
    p._p.append(wrapper[0])
    return p


def add_inline_equation_native(paragraph, latex: str, base_size: int = BODY_SIZE):
    """Append a native OMML inline math run to an existing paragraph."""
    xml_str = _latex_to_omml_xml(latex, display=False)
    wrapper = parse_xml(f"<m:tmp {_MATH_NSDECLS}>{xml_str}</m:tmp>")
    paragraph._p.append(wrapper[0])


def add_formatted_runs_with_math(paragraph, text: str, base_size=BODY_SIZE):
    """Like add_formatted_runs, but also converts $...$ inline LaTeX spans into
    native OMML math objects instead of plain text."""
    # Split on inline math first ($...$, not $$...$$), then run the normal
    # bold/italic/code formatter on the surrounding text segments.
    math_pattern = re.compile(r"\$([^$\n]+)\$")
    pos = 0
    for m in math_pattern.finditer(text):
        if m.start() > pos:
            add_formatted_runs(paragraph, text[pos:m.start()], base_size=base_size)
        add_inline_equation_native(paragraph, m.group(1), base_size=base_size)
        pos = m.end()
    if pos < len(text):
        add_formatted_runs(paragraph, text[pos:], base_size=base_size)


def clean_heading(text: str) -> str:
    # MD may use "2.1 Title" or "2.1. Title"; strip either before renumbering.
    return re.sub(r"^\d+(\.\d+)*\.?\s+", "", text).strip()


def is_identity_front_matter(stripped: str) -> bool:
    starts = (
        "**Running head:**", "**Authors:**", "**Author details:**",
        "**Correspondence:**", "- Harsh Dagar", "- Gunjan Bhandari",
    )
    if any(stripped.startswith(s) for s in starts):
        return True
    # Plain author / affiliation lines under the title (before Abstract)
    if "Dagar" in stripped or "Bhandari" in stripped:
        return True
    if "orcid.org" in stripped.lower():
        return True
    if stripped.startswith("¹") or stripped.startswith("1 "):
        return True
    if "ICAR" in stripped and "Karnal" in stripped:
        return True
    return False


# --------------------------------------------------------------------------- title page
def build_title_page():
    doc = Document()
    configure_styles(doc, line_numbers=False)

    add_para(doc, "Title page - Energy Policy", size=11, italic=True)
    add_para(doc, TITLE, size=14, bold=True, center=True)
    add_para(doc, "")

    add_para(doc, "Harsh Dagar¹* and Gunjan Bhandari¹", size=12, bold=True, center=True)
    add_para(
        doc,
        "¹ Division of Dairy Economics, Statistics and Management, "
        "ICAR-National Dairy Research Institute (Deemed University), "
        "Karnal, Haryana 132001, India",
        size=11, center=True,
    )
    add_para(doc, "")

    add_para(doc, "Author details", bold=True)
    add_para(
        doc,
        "Harsh Dagar, Research Scholar; ORCID: https://orcid.org/0009-0008-7394-130X; "
        "Email: harshdagar5879@gmail.com",
        size=11,
    )
    add_para(
        doc,
        "Gunjan Bhandari, Scientist; ORCID: https://orcid.org/0000-0001-6004-7642; "
        "Email: gunjanbhandari5@gmail.com",
        size=11,
    )
    add_para(doc, "")

    add_para(doc, "Corresponding author", bold=True)
    add_para(
        doc,
        "Harsh Dagar, Division of Dairy Economics, Statistics and Management, "
        "ICAR-National Dairy Research Institute, Karnal, Haryana 132001, India. "
        "Email: harshdagar5879@gmail.com",
        size=11,
    )
    add_para(doc, "")

    add_para(doc, "Declarations of interest", bold=True)
    add_para(doc, "The authors declare no competing interests.", size=11)

    add_para(doc, "Funding", bold=True)
    add_para(
        doc,
        "This research did not receive any specific grant from funding agencies in the "
        "public, commercial, or not-for-profit sectors.",
        size=11,
    )

    add_para(doc, "Acknowledgements", bold=True)
    add_para(
        doc,
        "We thank colleagues in the Division of Dairy Economics, Statistics and Management "
        "at ICAR-National Dairy Research Institute, Karnal, for institutional support. "
        "Errors remain ours.",
        size=11,
    )

    add_para(doc, "Author contributions (CRediT)", bold=True)
    add_para(
        doc,
        "H.D.: Conceptualisation; Data curation; Formal analysis; Investigation; "
        "Methodology; Software; Validation; Visualisation; Writing (original draft); "
        "Writing (review and editing). G.B.: Supervision; Validation; "
        "Writing (review and editing). Both authors approved the final manuscript.",
        size=11,
    )

    add_para(doc, "Article type", bold=True)
    add_para(doc, "Full Length Article. Target journal: Energy Policy (Elsevier).", size=11)

    path = OUTDIR / "01_Title_Page.docx"
    doc.save(path)
    print(f"Wrote {path}")


# --------------------------------------------------------------------------- highlights
def build_highlights():
    doc = Document()
    configure_styles(doc, line_numbers=False)
    add_para(doc, "Highlights", size=14, bold=True)
    add_para(
        doc,
        "Energy Policy requires 3-5 highlights, each a maximum of 85 characters including spaces.",
        size=10, italic=True,
    )
    for b in HIGHLIGHTS:
        n = len(b)
        if n > 85:
            raise ValueError(f"Highlight too long ({n}): {b}")
        p = doc.add_paragraph(style="List Bullet")
        set_double(p)
        run = p.add_run(b)
        set_run_font(run, size=BODY_SIZE)
    path = OUTDIR / "02_Highlights.docx"
    doc.save(path)
    print(f"Wrote {path}")


# --------------------------------------------------------------------------- manuscript (no authors)
def build_manuscript():
    text = MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_styles(doc, line_numbers=True)

    i = 0
    title_done = False
    sec = 0
    sub = 0
    subsub = 0
    in_refs = False
    skip_section = False

    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped or stripped == "---":
            i += 1
            continue

        # Title only — no authors
        if stripped.startswith("# ") and not title_done:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_double(p)
            run = p.add_run(stripped[2:].strip())
            set_run_font(run, size=14, bold=True)
            title_done = True
            i += 1
            continue

        # Drop all identity front matter
        if is_identity_front_matter(stripped):
            i += 1
            continue

        if stripped.startswith("**Keywords:**"):
            p = doc.add_paragraph()
            set_double(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_formatted_runs(p, stripped, base_size=11)
            for r in p.runs:
                set_run_font(r, size=11, italic=True, bold=r.bold)
            i += 1
            continue

        if stripped.startswith("## "):
            h = clean_heading(stripped[3:])
            key = h.upper()

            # End skip when a new section starts
            skip_section = key in STRIP_SECTIONS
            if skip_section:
                i += 1
                continue

            if key in UNNUMBERED or key.startswith("SUPPLEMENTARY"):
                if key == "REFERENCES":
                    in_refs = True
                display = "Abstract" if key == "ABSTRACT" else h
                p = doc.add_heading(display, level=1)
                force_heading_font(p, 14)
                if key == "ABSTRACT":
                    sec = 0
                i += 1
                continue

            sec += 1
            sub = 0
            p = doc.add_heading(f"{sec}. {h}", level=1)
            force_heading_font(p, 14)
            i += 1
            continue

        if skip_section:
            i += 1
            continue

        if stripped.startswith("#### "):
            h = clean_heading(stripped[5:])
            label = f"{sec}.{sub}.{subsub + 1} {h}" if sec > 0 else h
            subsub += 1
            p = doc.add_heading(label, level=3)
            force_heading_font(p, 11)
            i += 1
            continue

        if stripped.startswith("### "):
            h = clean_heading(stripped[4:])
            if sec > 0:
                sub += 1
                label = f"{sec}.{sub} {h}"
            else:
                label = h
            subsub = 0
            p = doc.add_heading(label, level=2)
            force_heading_font(p, 12)
            i += 1
            continue

        # Main-text figures pulled from the SI executive set
        if stripped.startswith("**FIGURE_EMBED:**"):
            payload = stripped.replace("**FIGURE_EMBED:**", "", 1).strip()
            if "|" not in payload:
                raise ValueError(f"Bad FIGURE_EMBED line: {stripped}")
            stem, caption = [x.strip() for x in payload.split("|", 1)]
            add_figure(doc, stem, caption)
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_md_table(lines, i)
            add_table(doc, rows)
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            set_double(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_formatted_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            for r in p.runs:
                set_run_font(r, size=BODY_SIZE, bold=r.bold, italic=r.italic)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_double(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_formatted_runs(p, stripped[2:])
            for r in p.runs:
                set_run_font(r, size=BODY_SIZE, bold=r.bold, italic=r.italic)
            i += 1
            continue

        if is_equation(stripped) or stripped.startswith("$$"):
            if stripped.startswith("$$") and not stripped.endswith("$$"):
                bits = [stripped]
                i += 1
                while i < len(lines) and not lines[i].strip().endswith("$$"):
                    bits.append(lines[i].strip())
                    i += 1
                if i < len(lines):
                    bits.append(lines[i].strip())
                eq = " ".join(bits)
                i += 1
            else:
                eq = stripped
                i += 1
            eq_latex = eq.strip()
            if eq_latex.startswith("$$"):
                eq_latex = eq_latex[2:]
            if eq_latex.endswith("$$"):
                eq_latex = eq_latex[:-2]
            add_block_equation_native(doc, eq_latex.strip())
            continue

        p = doc.add_paragraph()
        set_double(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if in_refs:
            p.paragraph_format.first_line_indent = Cm(-0.75)
            p.paragraph_format.left_indent = Cm(0.75)
        elif not stripped.startswith("**"):
            p.paragraph_format.first_line_indent = Cm(1.27)
        add_formatted_runs_with_math(p, stripped, base_size=BODY_SIZE)
        i += 1

    path = OUTDIR / "03_Manuscript.docx"
    doc.save(path)
    print(f"Wrote {path} ({path.stat().st_size/1024:.1f} KB)")


# --------------------------------------------------------------------------- cover letter
def build_cover_letter():
    doc = Document()
    configure_styles(doc, line_numbers=False)

    add_para(doc, "Harsh Dagar", bold=True)
    add_para(doc, "Division of Dairy Economics, Statistics and Management")
    add_para(doc, "ICAR-National Dairy Research Institute")
    add_para(doc, "Karnal, Haryana 132001, India")
    add_para(doc, "Email: harshdagar5879@gmail.com")
    add_para(doc, "ORCID: https://orcid.org/0009-0008-7394-130X")
    add_para(doc, "")
    add_para(doc, "Editor-in-Chief")
    add_para(doc, "Energy Policy")
    add_para(doc, "")
    add_para(doc, "Dear Editor,")
    add_para(doc, "")
    add_para(
        doc,
        f'Please find enclosed our manuscript, "{TITLE}," which we submit as a Full Length '
        "Article for consideration in Energy Policy. The authors are Harsh Dagar "
        "(corresponding) and Gunjan Bhandari.",
        indent=True,
    )
    add_para(
        doc,
        "The paper asks whether subsidised solar pumps under India's PM-KUSUM scheme deliver "
        "unconditional carbon gains or partly shift costs onto groundwater. Using a state-year "
        "panel, associational two-way fixed effects, and an instrumental-variables design with "
        "placebos and stress tests, we find a positive intensive-margin groundwater response "
        "concentrated in Component B (standalone diesel-replacement pumps). Gross carbon "
        "abatement is about 1.93 Mt CO2 per year; net abatement is about 1.65 Mt after an "
        "intensive-margin rebound of about 0.28 Mt. The rebound share rises from about 11% in "
        "safe aquifers to about 31% in over-exploited ones because deeper water tables cost more "
        "energy to lift. We recommend conditioning Component B and haircutting climate credits "
        "where aquifers are stressed, rather than halting solar irrigation, and treating "
        "Components B and C as separate instruments.",
        indent=True,
    )
    add_para(
        doc,
        "We believe the manuscript fits Energy Policy because it links a major national energy "
        "programme to clear implications for MNRE, state agencies and results-based climate "
        "finance. The main text embeds four figures; Tables S1-S3 and Figures S1-S10 are in the "
        "Supplementary Information. The manuscript is original, not under review elsewhere, and "
        "approved by both authors.",
        indent=True,
    )
    add_para(doc, "Thank you for your consideration.", indent=True)
    add_para(doc, "")
    add_para(doc, "Yours sincerely,")
    add_para(doc, "Harsh Dagar")
    add_para(doc, "On behalf of Harsh Dagar and Gunjan Bhandari")

    path = OUTDIR / "04_Cover_Letter.docx"
    doc.save(path)
    print(f"Wrote {path}")


def sync_si_and_figures():
    """Copy SI DOCX + executive visuals + SI table CSVs into the READY folder."""
    import shutil

    # Rebuild SI DOCX from current script if present
    si_builder = ROOT / "scripts" / "build_si_docx.py"
    if si_builder.exists():
        import runpy
        runpy.run_path(str(si_builder), run_name="__main__")

    si_src = ROOT / "outputs" / "manuscript_PAPER3_SupplementaryInformation.docx"
    si_dst = OUTDIR / "05_Supplementary_Information.docx"
    if not si_src.exists():
        raise FileNotFoundError(si_src)
    shutil.copy2(si_src, si_dst)
    print(f"Wrote {si_dst} ({si_dst.stat().st_size/1024:.1f} KB)")

    fig_src = ROOT / "outputs" / "figures" / "executive"
    fig_dst = OUTDIR / "figures"
    fig_dst.mkdir(parents=True, exist_ok=True)
    n_png = n_pdf = 0
    for ext in ("*.png", "*.pdf"):
        for f in sorted(fig_src.glob(ext)):
            shutil.copy2(f, fig_dst / f.name)
            if f.suffix.lower() == ".png":
                n_png += 1
            else:
                n_pdf += 1
    print(f"Copied {n_png} PNG + {n_pdf} PDF figures -> {fig_dst}")

    tab_dst = OUTDIR / "tables"
    tab_dst.mkdir(parents=True, exist_ok=True)
    for name in (
        "SI_S1_carbon_sensitivity_matrix.csv",
        "SI_S2_iv_diagnostics.csv",
        "SI_S3_social_cost_ledger.csv",
        "executive_visual_inventory.csv",
    ):
        src = ROOT / "outputs" / "tables" / name
        if src.exists():
            shutil.copy2(src, tab_dst / name)
    print(f"Copied SI table CSVs -> {tab_dst}")


def write_readme():
    text = """Energy Policy - upload these separate files:

1. 01_Title_Page.docx                 - title, authors, affiliations, ORCID, CRediT, declarations
2. 02_Highlights.docx                 - 5 highlights (each <= 85 characters)
3. 03_Manuscript.docx                 - anonymised main text (NO author block); Times New Roman; justified; double-spaced; continuous line numbers; Figures 1-4 embedded
4. 04_Cover_Letter.docx               - submission letter
5. 05_Supplementary_Information.docx  - SI Tables S1-S3 + Figures S1-S10 (embedded)
6. figures/                           - same visuals as PNG + PDF @ 400 dpi (optional separate upload)
7. tables/                            - SI_S1 / SI_S2 / SI_S3 CSV source files
8. manuscript_source.md               - master markdown text (not for upload; edit this, then rebuild)

Notes:
- Do NOT put affiliations inside 03_Manuscript.docx (already stripped).
- Main-text figures: V2, V4, V5, V9 (geography; paradox scales; component meters; carbon balance).
- Rebuild anytime:  python scripts/build_energy_policy_pack.py
- Portal: https://www.editorialmanager.com/jepo/
"""
    path = OUTDIR / "00_UPLOAD_THESE.txt"
    path.write_text(text, encoding="utf-8")
    print(f"Wrote {path}")


def main():
    build_title_page()
    build_highlights()
    build_manuscript()
    build_cover_letter()
    sync_si_and_figures()
    write_readme()
    # Keep a reference copy of the master markdown alongside the pack
    import shutil
    shutil.copy2(MD, OUTDIR / "manuscript_source.md")
    print(f"Also copied -> {OUTDIR / 'manuscript_source.md'}")


if __name__ == "__main__":
    main()
