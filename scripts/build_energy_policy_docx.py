"""
Build Energy Policy submission DOCX from the Paper 3 manuscript markdown.

Energy Policy technical-screening conventions applied:
- Times New Roman, 12 pt body
- Double-line spacing throughout
- Single-column layout, 1-inch margins
- Numbered section/subsection headings (1, 1.1, ...)
- Editable Word tables
- Highlights as a separate DOCX (mandatory for Energy Policy)
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "outputs" / "manuscript_PAPER3_NatureSustainability.md"
OUT = ROOT / "outputs" / "manuscript_PAPER3_EnergyPolicy.docx"
OUT_FALLBACK = ROOT / "outputs" / "manuscript_PAPER3_EnergyPolicy_NEW.docx"
HIGHLIGHTS_OUT = ROOT / "outputs" / "Highlights_EnergyPolicy.docx"

BODY_SIZE = 12
TABLE_SIZE = 10
META_SIZE = 11

UNNUMBERED = {
    "ABSTRACT", "METHODS SUMMARY", "DATA AVAILABILITY", "CODE AVAILABILITY",
    "ACKNOWLEDGEMENTS", "AUTHOR CONTRIBUTIONS", "COMPETING INTERESTS",
    "SUPPLEMENTARY INFORMATION", "REFERENCES", "HIGHLIGHTS",
}


def set_run_font(run, name="Times New Roman", size=BODY_SIZE, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=10)


def set_double(paragraph):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(BODY_SIZE)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    for i, size in [(1, 14), (2, 12), (3, 12)]:
        style = doc.styles[f"Heading {i}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(0)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(p)


def add_formatted_runs(paragraph, text: str, base_size=BODY_SIZE):
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Courier New", size=base_size - 1)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=base_size, italic=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "666666")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.autofit = True
    for i, row_data in enumerate(rows):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            val = row_data[j] if j < len(row_data) else ""
            val = re.sub(r"\*\*([^*]+)\*\*", r"\1", val)
            run = p.add_run(val.strip())
            set_run_font(run, size=TABLE_SIZE, bold=(i == 0))
            set_cell_border(cell)
        if i == 0:
            for cell in table.rows[0].cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "E7E6E6")
                shd.set(qn("w:val"), "clear")
                tcPr.append(shd)
    spacer = doc.add_paragraph()
    set_double(spacer)


def parse_md_table(lines: list[str], start: int):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.match(r"^\|[\s:\-|]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def is_equation(line: str) -> bool:
    s = line.strip()
    return s.startswith("$$") or (s.startswith("$") and s.endswith("$") and len(s) > 2)


def clean_heading(text: str) -> str:
    return re.sub(r"^\d+(\.\d+)*\.\s+", "", text).strip()


def build_highlights_docx():
    bullets = [
        "KUSUM intensity raises GW stage by +0.10 pp (IV; F = 20.6).",
        "Component B drives the effect (F = 54.3); Component C is weak.",
        "Carbon rebound offsets ~32% of gains in over-exploited zones.",
        "IV sign survives dropping Punjab, Haryana and Rajasthan.",
        "Condition Component B and haircut credits; do not halt pumps.",
    ]
    doc = Document()
    configure_styles(doc)
    p = doc.add_paragraph()
    set_double(p)
    run = p.add_run("Highlights")
    set_run_font(run, size=14, bold=True)
    for b in bullets:
        n = len(b)
        if n > 85:
            raise ValueError(f"Highlight too long ({n} chars): {b}")
        p = doc.add_paragraph(style="List Bullet")
        set_double(p)
        run = p.add_run(b)
        set_run_font(run, size=BODY_SIZE)
    doc.save(HIGHLIGHTS_OUT)
    print(f"Wrote {HIGHLIGHTS_OUT}")


def is_front_meta(stripped: str, sec: int) -> bool:
    if sec > 0:
        return False
    starts = (
        "**Running head:**", "**Authors:**", "**Author details:**",
        "**Correspondence:**", "- Harsh Dagar", "- Gunjan Bhandari",
    )
    if any(stripped.startswith(s) for s in starts):
        return True
    if stripped.startswith("¹ ") or stripped.startswith("1 "):
        return True
    return False


def build():
    text = MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_styles(doc)

    i = 0
    title_done = False
    sec = 0
    sub = 0
    in_refs = False
    numbering_on = False

    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        # Title
        if stripped.startswith("# ") and not title_done:
            title = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_double(p)
            run = p.add_run(title)
            set_run_font(run, size=14, bold=True)
            title_done = True
            i += 1
            continue

        if is_front_meta(stripped, sec):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_double(p)
            add_formatted_runs(p, stripped, base_size=META_SIZE)
            i += 1
            continue

        if stripped.startswith("**Keywords:**"):
            p = doc.add_paragraph()
            set_double(p)
            add_formatted_runs(p, stripped, base_size=META_SIZE)
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # ## headings
        if stripped.startswith("## "):
            h = clean_heading(stripped[3:])
            key = h.upper()
            if key in UNNUMBERED or key.startswith("SUPPLEMENTARY"):
                if key == "REFERENCES":
                    in_refs = True
                    numbering_on = False
                display = "Abstract" if key == "ABSTRACT" else h
                p = doc.add_heading(display, level=1)
                set_double(p)
                # Start numbering after Abstract
                if key == "ABSTRACT":
                    numbering_on = True
                    sec = 0
                i += 1
                continue

            if numbering_on or True:
                # Number all main research sections; skip only explicit unnumbered set
                sec += 1
                sub = 0
                p = doc.add_heading(f"{sec}. {h}", level=1)
                set_double(p)
                numbering_on = True
            i += 1
            continue

        # ### subsections
        if stripped.startswith("### "):
            h = clean_heading(stripped[4:])
            if sec > 0:
                sub += 1
                label = f"{sec}.{sub} {h}"
            else:
                label = h
            p = doc.add_heading(label, level=2)
            set_double(p)
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(clean_heading(stripped[2:]), level=1)
            set_double(p)
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_md_table(lines, i)
            add_table(doc, rows)
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            set_double(p)
            add_formatted_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_double(p)
            add_formatted_runs(p, stripped[2:])
            i += 1
            continue

        if is_equation(stripped) or stripped.startswith("$$"):
            if stripped.startswith("$$") and not stripped.endswith("$$"):
                bits = [stripped]
                i += 1
                while i < len(lines) and not lines[i].strip().endswith("$$"):
                    bits.append(lines[i].strip())
                    i += 1
                if i < len(lines):
                    bits.append(lines[i].strip())
                eq = " ".join(bits)
                i += 1
            else:
                eq = stripped
                i += 1
            eq_clean = eq.replace("$$", "").replace("$", "").strip()
            eq_clean = (eq_clean
                        .replace(r"\cdot", "·")
                        .replace(r"\log", "log")
                        .replace(r"\beta", "β")
                        .replace(r"\gamma", "γ")
                        .replace(r"\alpha", "α")
                        .replace(r"\lambda", "λ")
                        .replace(r"\varepsilon", "ε")
                        .replace(r"\text{", "")
                        .replace("}", "")
                        .replace("{", "")
                        .replace(r"\widehat", "")
                        .replace(r"\hat", "")
                        .replace(r"\Delta", "Δ"))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_double(p)
            run = p.add_run(eq_clean)
            set_run_font(run, size=BODY_SIZE, italic=True)
            continue

        p = doc.add_paragraph()
        set_double(p)
        if in_refs:
            p.paragraph_format.first_line_indent = Cm(-0.75)
            p.paragraph_format.left_indent = Cm(0.75)
        elif not stripped.startswith("**"):
            p.paragraph_format.first_line_indent = Cm(1.27)
        add_formatted_runs(p, stripped, base_size=BODY_SIZE)
        i += 1

    try:
        doc.save(OUT)
        print(f"Wrote {OUT} ({OUT.stat().st_size/1024:.1f} KB)")
    except PermissionError:
        doc.save(OUT_FALLBACK)
        print(f"Locked; wrote {OUT_FALLBACK}")


if __name__ == "__main__":
    build_highlights_docx()
    build()
