import re
from docx import Document

text = open('scripts/_manuscript_docx_text_v2.txt', encoding='utf-8').read()

m = re.search(r'1\. INTRODUCTION\n(.+?)\nMETHODS SUMMARY', text, re.DOTALL)
body = m.group(1)
words = re.findall(r'\b[A-Za-z0-9]+\b', body)
print('Body prose word count:', len(words))

m2 = re.search(r'Keywords:\s*(.+)', text)
print('Keywords:', m2.group(1))

doc = Document('FINAL_SUBMISSION_PACKAGE/03_Manuscript.docx')
print('Tables:', len(doc.tables))
print('Paragraphs:', len(doc.paragraphs))
xml = doc.element.xml
print('oMath count:', xml.count('<m:oMath>'))
print('oMathPara opens:', xml.count('<m:oMathPara>'))

# check intro paragraph reads sensibly
for p in doc.paragraphs:
    if p.text.startswith('Solar irrigation is considered'):
        print()
        print('INTRO paragraph (fixed):')
        print(p.text[:500])
        break
