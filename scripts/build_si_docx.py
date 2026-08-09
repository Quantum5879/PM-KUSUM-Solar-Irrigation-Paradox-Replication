"""
Build Supplementary Information DOCX for Paper 3.
Sequence: Tables S1-S3, then Figures S1-S10 (one by one).
Does NOT alter the main manuscript body beyond the separate SI file.
"""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
TAB = ROOT / "outputs" / "tables"
FIG = ROOT / "outputs" / "figures" / "executive"
OUT = ROOT / "outputs" / "manuscript_PAPER3_SupplementaryInformation.docx"

# Figure order: S1..S12. V2, V4, V5 and V9 are Manuscript Figures 1-4 and are NOT
# repeated here to avoid duplicating content already in the main text.
FIGURES = [
    ("S1", "V1_causal_storyboard", "Causal pathway"),
    ("S2", "V3_event_ribbon_timeline", "Event study"),
    ("S3", "V6_cate_ridges_by_stress", "Heterogeneous effects"),
    ("S4", "V7_rebound_heatmap", "Rebound sensitivity"),
    ("S5", "V8_loo_lollipop_spine", "Leave-one-out"),
    ("S6", "V10_governance_matrix", "Governance matrix"),
    ("S7", "V11_treatment_balance", "Treatment balance"),
    ("S8", "V12_assumption_ledger", "Assumption ledger"),
    ("S9", "V13_confirmatory_forest", "Confirmatory battery"),
    ("S10", "V14_yearly_reduced_form", "Year-by-year reduced form"),
    ("S11", "V15_csa_event_study", "Callaway-Sant'Anna event study"),
    ("S12", "V16_ccts_revenue_fanchart", "CCTS revenue by state"),
]


def set_run_font(run, name="Times New Roman", size=11, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(8)

    for i, size in [(1, 16), (2, 13), (3, 12)]:
        style = doc.styles[f"Heading {i}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.space_before = Pt(14 if i == 1 else 10)
        style.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(p)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "666666")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_df_table(doc: Document, df: pd.DataFrame, float_cols=None, max_rows=None):
    float_cols = float_cols or {}
    view = df.copy()
    if max_rows is not None and len(view) > max_rows:
        view = view.head(max_rows)
    for col, nd in float_cols.items():
        if col in view.columns:
            view[col] = view[col].map(
                lambda x, n=nd: "" if pd.isna(x) else (f"{x:.{n}f}" if isinstance(x, (int, float)) else str(x))
            )
    # stringify rest
    for col in view.columns:
        if col not in float_cols:
            view[col] = view[col].map(lambda x: "" if pd.isna(x) else str(x))

    headers = [str(c) for c in view.columns]
    rows = [headers] + view.values.tolist()
    table = doc.add_table(rows=len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            run = p.add_run(str(val))
            set_run_font(run, size=8, bold=(i == 0))
            set_cell_border(cell)
    doc.add_paragraph()


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=10, italic=True)


def add_heading_para(doc: Document, text: str, level=1):
    doc.add_heading(text, level=level)


def page_break(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def build():
    doc = Document()
    configure_styles(doc)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Supplementary Information")
    set_run_font(run, size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Do Subsidised Solar Pumps Really Cut Carbon or Shift the Cost to Groundwater?\n"
        "Evidence from India's PM-KUSUM"
    )
    set_run_font(run, size=12, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Harsh Dagar and Gunjan Bhandari\n"
        "Division of Dairy Economics, Statistics and Management\n"
        "ICAR-National Dairy Research Institute (Deemed University), Karnal, Haryana 132001, India\n"
        "ORCID: H.D. https://orcid.org/0009-0008-7394-130X ; "
        "G.B. https://orcid.org/0000-0001-6004-7642"
    )
    set_run_font(run, size=11)

    p = doc.add_paragraph()
    run = p.add_run(
        "This Supplementary Information accompanies the main article. It reports every "
        "table and figure produced during the analysis that is not already shown in the "
        "main text. Figures already embedded as Figures 1-4 in the manuscript (geographic "
        "alignment, paradox scales, component identification, and carbon balance) are not "
        "repeated here. Supplementary Tables S1-S10 cover carbon sensitivity, IV diagnostics, "
        "the social-cost ledger, treatment balance, the assumption/threat ledger, the full "
        "confirmatory battery, the complete leave-one-out sweep, year-by-year reduced-form "
        "estimates, the full Callaway-Sant'Anna event-study path, and CCTS carbon-credit "
        "revenue by state. Supplementary Figures S1-S12 cover the causal pathway, event "
        "study, heterogeneous effects, rebound sensitivity, leave-one-out, the governance "
        "matrix, treatment balance, the assumption ledger, the confirmatory battery, the "
        "year-by-year reduced form, the Callaway-Sant'Anna event study, and CCTS revenue by "
        "state. Figures are rendered at 400 dpi."
    )
    set_run_font(run, size=10)

    # -------------------- TABLES --------------------
    add_heading_para(doc, "Supplementary Tables", level=1)

    # Table S1
    add_heading_para(doc, "Table S1. Carbon sensitivity matrix", level=2)
    add_caption(
        doc,
        "Gross, rebound and net abatement (Mt CO2/yr) across diesel displacement, "
        "pump energy intensity, lift depth and grid emission factor. "
        "is_central = 1 marks the paper's central bridge case "
        "(600 L, e = 0.004, lift = 20 m, EF = 0.7383).",
    )
    s1 = pd.read_csv(TAB / "SI_S1_carbon_sensitivity_matrix.csv")
    s1_show = s1.rename(columns={
        "fuel_litres_displaced": "Fuel (L)",
        "e_kWh_per_m3_m": "e (kWh/m3/m)",
        "lift_base_m": "Lift (m)",
        "ef_grid": "EF grid",
        "gross_Mt": "Gross Mt",
        "rebound_Mt": "Rebound Mt",
        "net_Mt": "Net Mt",
        "offset_pct": "Offset %",
        "is_central": "Central",
    })
    add_df_table(
        doc, s1_show,
        float_cols={
            "e (kWh/m3/m)": 4, "Lift (m)": 0, "EF grid": 4,
            "Gross Mt": 3, "Rebound Mt": 3, "Net Mt": 3, "Offset %": 1,
        },
    )

    page_break(doc)

    # Table S2
    add_heading_para(doc, "Table S2. IV diagnostics", level=2)
    add_caption(
        doc,
        "Instrumental-variable specifications for groundwater stage and related outcomes. "
        "Weak_IV_flag marks first-stage F < 10. The final row summarises leave-one-out sign "
        "stability across all specifications rather than a single hypothesis test, so its "
        "p-value cell is blank.",
    )
    s2 = pd.read_csv(TAB / "SI_S2_iv_diagnostics.csv")
    s2_show = s2.rename(columns={
        "spec": "Specification",
        "outcome": "Outcome",
        "beta": "Beta",
        "se": "SE",
        "pvalue": "p-value",
        "first_stage_F": "First-stage F",
        "n": "N",
        "weak_IV_flag": "Weak IV",
    })
    add_df_table(
        doc, s2_show,
        float_cols={"Beta": 4, "SE": 4, "p-value": 4, "First-stage F": 2},
    )

    page_break(doc)

    # Table S3
    add_heading_para(doc, "Table S3. Social cost ledger", level=2)
    add_caption(
        doc,
        "Gross abatement value, rebound cost, net value and over-claim cost (INR crore) "
        "at alternative social cost of carbon (INR per tCO2). Zone rows use SCC = 1,500 INR/t. "
        "Offset share is only meaningful within a zone breakdown, so it is blank for the "
        "top four SCC-sensitivity rows.",
    )
    s3 = pd.read_csv(TAB / "SI_S3_social_cost_ledger.csv")
    s3_show = s3.rename(columns={
        "scc_INR_per_t": "SCC (INR/t)",
        "gross_value_INR_crore": "Gross value",
        "rebound_cost_INR_crore": "Rebound cost",
        "net_value_INR_crore": "Net value",
        "overclaim_cost_INR_crore": "Over-claim cost",
        "offset_share": "Offset share",
    })
    add_df_table(
        doc, s3_show,
        float_cols={
            "Gross value": 2, "Rebound cost": 2, "Net value": 2,
            "Over-claim cost": 2, "Offset share": 3,
        },
    )

    page_break(doc)

    # Table S4
    add_heading_para(doc, "Table S4. Treatment balance", level=2)
    add_caption(
        doc,
        "Mean covariate values for states above vs. below the top-tercile KUSUM intensity "
        "threshold in the latest panel year. Blank cells indicate the variable was not "
        "available for a sufficient number of states in that group.",
    )
    s4 = pd.read_csv(TAB / "SI_S4_treatment_balance.csv")
    add_df_table(
        doc, s4,
        float_cols={"High-KUSUM mean": 2, "Low-KUSUM mean": 2},
    )

    page_break(doc)

    # Table S5
    add_heading_para(doc, "Table S5. Assumption and threat-to-validity ledger", level=2)
    add_caption(
        doc,
        "Every identification assumption behind the design, the specific threat it faces, "
        "the test applied, and its resolution status (fixed, checked, partial, or rejected). "
        "Referenced in Section 4.7 of the manuscript.",
    )
    s5 = pd.read_csv(TAB / "SI_S5_assumption_ledger.csv")
    add_df_table(doc, s5)

    page_break(doc)

    # Table S6
    add_heading_para(doc, "Table S6. Confirmatory battery: component splits and placebos", level=2)
    add_caption(
        doc,
        "Fixed-effects component splits (Component B vs. C on tube-wells and groundwater "
        "stage), the continuous stress-by-DISCOM interaction model, the non-interpolated "
        "groundwater subsample, and the canal-area placebo. Backs the numbers cited in "
        "Sections 5.5 and 5.7 of the manuscript.",
    )
    s6 = pd.read_csv(TAB / "SI_S6_confirmatory_battery.csv")
    add_df_table(
        doc, s6,
        float_cols={"Beta": 4, "SE": 4, "p-value": 4},
    )

    page_break(doc)

    # Table S7
    add_heading_para(doc, "Table S7. Leave-one-out IV, all states", level=2)
    add_caption(
        doc,
        "Baseline instrumental-variables coefficient re-estimated after dropping each state "
        "in turn, for all 34 states with usable variation. Figure S5 (leave-one-out) plots "
        "a subset of these rows; this table reports the complete sweep.",
    )
    s7 = pd.read_csv(TAB / "SI_S7_leave_one_out_full.csv")
    add_df_table(
        doc, s7,
        float_cols={"Beta": 4, "SE": 4, "p-value": 4, "First-stage F": 2},
    )

    page_break(doc)

    # Table S8
    add_heading_para(doc, "Table S8. Year-by-year reduced-form estimates", level=2)
    add_caption(
        doc,
        "Cross-sectional reduced-form regression of groundwater stage on the instrument, "
        "estimated separately for each fiscal year from 2019 onward. Backs the timing "
        "argument in Section 5.8 of the manuscript.",
    )
    s8 = pd.read_csv(TAB / "SI_S8_year_by_year_reduced_form.csv")
    add_df_table(
        doc, s8,
        float_cols={"Beta (reduced form)": 5, "SE": 5, "p-value": 4},
    )

    page_break(doc)

    # Table S9
    add_heading_para(doc, "Table S9. Callaway-Sant'Anna event study, full path", level=2)
    add_caption(
        doc,
        "Group-time average treatment effect on tube-well area at each year relative to "
        "KUSUM onset, with pointwise 95 percent confidence bands. Backs the -506.9 ha claim "
        "at t = +3 cited in Section 5.3 of the manuscript.",
    )
    s9 = pd.read_csv(TAB / "SI_S9_callaway_santanna_full.csv")
    add_df_table(
        doc, s9,
        float_cols={"ATT (ha)": 1, "SE": 1, "95% CI lower": 1, "95% CI upper": 1},
    )

    page_break(doc)

    # Table S10
    add_heading_para(doc, "Table S10. CCTS carbon-credit revenue by state", level=2)
    add_caption(
        doc,
        "Monte Carlo valuation (10,000 draws, triangular carbon-price distribution) of net "
        "abatement at each state's current carbon credit trading scheme (CCTS) exposure. "
        "States with zero estimated net abatement are omitted. Supporting result; not a "
        "causal estimate.",
    )
    s10 = pd.read_csv(TAB / "SI_S10_ccts_revenue_by_state.csv")
    add_df_table(
        doc, s10,
        float_cols={
            "Net abatement (tCO2/yr)": 1, "Revenue p5 (INR crore)": 2,
            "Revenue p50 (INR crore)": 2, "Revenue p95 (INR crore)": 2,
        },
    )

    # -------------------- FIGURES --------------------
    page_break(doc)
    add_heading_para(doc, "Supplementary Figures", level=1)

    for i, (sid, stem, title) in enumerate(FIGURES):
        if i > 0:
            page_break(doc)
        add_heading_para(doc, f"Figure {sid}. {title}", level=2)
        png = FIG / f"{stem}.png"
        if not png.exists():
            p = doc.add_paragraph()
            run = p.add_run(f"[Missing file: {png.name}]")
            set_run_font(run, size=10, italic=True)
            continue
        # Width ~6.3 inches for portrait readability
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(png), width=Inches(6.3))
        add_caption(doc, f"Figure {sid}. {title}. Source: {stem}.png (400 dpi).")

    try:
        doc.save(OUT)
        print(f"Wrote {OUT} ({OUT.stat().st_size/1024:.1f} KB)")
    except PermissionError:
        alt = OUT.with_name(OUT.stem + "_NEW.docx")
        doc.save(alt)
        print(f"SI DOCX locked; wrote {alt} ({alt.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    build()
