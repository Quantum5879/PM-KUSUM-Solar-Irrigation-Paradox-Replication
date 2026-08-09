"""
Build a properly formatted academic DOCX from the Paper 3 manuscript markdown.
- Real Word tables (not pipe-text)
- Journal-like typography (Times New Roman body, clear headings)
- Bold/italic runs, page numbers, readable margins
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "outputs" / "manuscript_PAPER3_NatureSustainability.md"
OUT = ROOT / "outputs" / "manuscript_PAPER3_NatureSustainability.docx"
OUT_FALLBACK = ROOT / "outputs" / "manuscript_PAPER3_NatureSustainability_FINAL.docx"


def set_run_font(run, name="Times New Roman", size=11, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(8)
    pf.space_before = Pt(0)

    for i, size in [(1, 16), (2, 13), (3, 11)]:
        style = doc.styles[f"Heading {i}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.space_before = Pt(14 if i == 1 else 10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(p)


def add_formatted_runs(paragraph, text: str, base_size=11):
    """Parse **bold**, *italic*, and `code` segments."""
    # Protect ** and * carefully
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=base_size - 1)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=base_size, italic=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size)


def shade_header_row(row):
    for cell in row.cells:
        tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9E2F3")
        shd.set(qn("w:val"), "clear")
        tc.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "666666")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.autofit = True
    for i, row_data in enumerate(rows):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            val = row_data[j] if j < len(row_data) else ""
            val = re.sub(r"\*\*([^*]+)\*\*", r"\1", val)
            run = p.add_run(val.strip())
            set_run_font(run, size=9, bold=(i == 0))
            set_cell_border(cell)
        if i == 0:
            for cell in table.rows[0].cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "D6E3F0")
                shd.set(qn("w:val"), "clear")
                tcPr.append(shd)
    doc.add_paragraph()


def parse_md_table(lines: list[str], start: int):
    """Return (rows, next_index)."""
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.match(r"^\|[\s:\-|]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def is_equation(line: str) -> bool:
    s = line.strip()
    return s.startswith("$$") or (s.startswith("$") and s.endswith("$") and len(s) > 2)


def build():
    text = MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_styles(doc)

    i = 0
    title_done = False
    while i < len(lines):
        line = lines[i]
        raw = line.rstrip()
        stripped = raw.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        # Title
        if stripped.startswith("# ") and not title_done:
            title = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(title)
            set_run_font(run, size=16, bold=True)
            title_done = True
            i += 1
            continue

        # Meta lines under title
        if stripped.startswith("**Running head:**") or stripped.startswith("**Authors:**") or stripped.startswith("**Correspondence:**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            add_formatted_runs(p, stripped, base_size=10)
            i += 1
            continue

        if stripped.startswith("**Keywords:**"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            add_formatted_runs(p, stripped, base_size=10)
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        # Tables
        if stripped.startswith("|"):
            rows, i = parse_md_table(lines, i)
            add_table(doc, rows)
            continue

        # Numbered / bullet lists
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, stripped[2:])
            i += 1
            continue

        # Equations
        if is_equation(stripped) or stripped.startswith("$$"):
            # gather multi-line $$ blocks
            eq = stripped
            if stripped.startswith("$$") and not stripped.endswith("$$"):
                i += 1
                bits = [stripped]
                while i < len(lines) and not lines[i].strip().endswith("$$"):
                    bits.append(lines[i].strip())
                    i += 1
                if i < len(lines):
                    bits.append(lines[i].strip())
                eq = " ".join(bits)
                i += 1
            else:
                i += 1
            eq_clean = eq.replace("$$", "").replace("$", "").strip()
            # light LaTeX cleanup for Word readability
            eq_clean = (eq_clean
                        .replace(r"\cdot", "·")
                        .replace(r"\log", "log")
                        .replace(r"\beta", "β")
                        .replace(r"\gamma", "γ")
                        .replace(r"\alpha", "α")
                        .replace(r"\lambda", "λ")
                        .replace(r"\varepsilon", "ε")
                        .replace(r"\text{", "")
                        .replace("}", "")
                        .replace("{", "")
                        .replace(r"\widehat", "")
                        .replace(r"\hat", "")
                        .replace(r"\Delta", "Δ"))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(eq_clean)
            set_run_font(run, size=11, italic=True)
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.5) if not stripped.startswith("**") else Cm(0)
        add_formatted_runs(p, stripped)
        i += 1

    try:
        doc.save(OUT)
        print(f"Wrote {OUT} ({OUT.stat().st_size/1024:.1f} KB)")
    except PermissionError:
        doc.save(OUT_FALLBACK)
        print(f"Main DOCX locked; wrote {OUT_FALLBACK} ({OUT_FALLBACK.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    build()
